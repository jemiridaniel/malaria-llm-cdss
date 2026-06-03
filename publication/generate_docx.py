"""
generate_docx.py
================
Converts publication/final_paper.md to publication/final_paper.docx
using python-docx for journal submission.

Pipeline:
  1. Parse markdown line-by-line, tracking section state
     (front_matter → abstract → body → references)
  2. Apply per-section formatting rules
  3. Inject actual figure images at the correct paragraph positions
  4. Build tables from pipe-delimited rows
  5. Apply academic Word styling (A4, 1-inch margins, header/footer)

Usage:
    python publication/generate_docx.py

Requirements:
    python-docx  (pip install python-docx)
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ─────────────────────────────────────────────────────────────────────
PUB_DIR = Path(__file__).parent
FIG_DIR = PUB_DIR / "figures"
MD_SRC  = PUB_DIR / "final_paper.md"
OUT_DOC = PUB_DIR / "final_paper.docx"

BODY_FONT = "Times New Roman"
CODE_FONT = "Courier New"
BODY_SIZE = 12

# ── Figure injection map ──────────────────────────────────────────────────────
# (trigger sentence in body text, figure filename, caption)
FIGURE_TRIGGERS = [
    (
        "Figure 1 illustrates the full system pipeline.",
        "fig1_system_architecture.png",
        "Figure 1. MalariaLLM System Architecture. Five-stage pipeline from Patient Input "
        "through the Rule-Based Classification Engine, LLM Reasoning Module, Structured "
        "Clinical Output, and PDF Report. The LLM receives the rule-derived severity "
        "classification and cannot alter it.",
    ),
    (
        "Figure 2 visualizes accuracy across all stages for both systems.",
        "fig2_accuracy_comparison.png",
        "Figure 2. Diagnostic Accuracy: Baseline vs. LLM-Enhanced Hybrid System. "
        "Side-by-side bars for No Malaria, Stage I, Stage II, Critical, and Overall "
        "(n\u2009=\u20091,682). pp\u00a0= percentage points.",
    ),
    (
        "Row-normalised confusion matrices for both systems are presented in Figure 3.",
        "fig3_confusion_matrices.png",
        "Figure 3. Confusion Matrices: Row-Normalised per True Stage (n\u2009=\u20091,682). "
        "Left: Baseline (Rule-Based Only). Right: LLM-Enhanced Hybrid. Cell values show "
        "row-normalised percentage and raw case count.",
    ),
    (
        "Figure 4 presents the processing time distribution and the "
        "accuracy\u2013latency trade-off.",
        "fig4_processing_time.png",
        "Figure 4. Processing Efficiency. Left: LLM inference time distribution with mean "
        "(6.3\u2009s) and median (5.4\u2009s) reference lines. Right: "
        "Accuracy\u2013latency trade-off: baseline (<0.01\u2009s, 30.62%) vs. hybrid "
        "(6.3\u2009s, 87.22%).",
    ),
]


# ── XML helpers ───────────────────────────────────────────────────────────────

def _shd(fill: str) -> OxmlElement:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), fill)
    return el


def set_cell_bg(cell, hex_color: str):
    cell._tc.get_or_add_tcPr().append(_shd(hex_color))


def set_para_shading(para, hex_color: str):
    para._p.get_or_add_pPr().append(_shd(hex_color))


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    bdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "AAAAAA")
        bdr.append(b)
    tblPr.append(bdr)


def add_para_border_bottom(para, color="CBD5E1"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def add_page_number_footer(doc):
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    for tag, attr_val in [("begin", None), ("instrText", "PAGE"), ("end", None)]:
        if tag == "instrText":
            el = OxmlElement("w:instrText")
            el.text = attr_val
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
        run._r.append(el)
    run.font.size = Pt(10)
    run.font.name = BODY_FONT


def add_doc_header(doc, text: str):
    header = doc.sections[0].header
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run.font.name = BODY_FONT


# ── Inline markdown parser ────────────────────────────────────────────────────

def parse_inline(para, text: str, size=BODY_SIZE, font=BODY_FONT,
                 bold=False, italic=False, color: RGBColor | None = None):
    """Add styled runs to *para* by parsing inline markdown tokens."""
    # Flatten markdown links [label](url) → label
    text = re.sub(r'\[([^\]]*)\]\([^)]*\)', r'\1', text)

    TOKEN = re.compile(r'(\*\*\*[^*]+?\*\*\*|\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`)')
    for part in TOKEN.split(text):
        if not part:
            continue
        run = para.add_run()
        run.font.name = font
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        if part.startswith("***") and part.endswith("***"):
            run.text = part[3:-3]
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
            run.italic = italic
        elif part.startswith("*") and part.endswith("*") and len(part) >= 3:
            run.text = part[1:-1]
            run.bold = bold
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) >= 3:
            run.text = part[1:-1]
            run.font.name = CODE_FONT
            run.font.size = Pt(size - 1)
        else:
            run.text = part
            run.bold = bold
            run.italic = italic


# ── Block builders ────────────────────────────────────────────────────────────

def new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6,
             lsp=1.5, left_indent=None, right_indent=None,
             first_line=None) -> object:
    """Add a paragraph with common formatting shortcuts."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = lsp
    if left_indent  is not None: pf.left_indent        = left_indent
    if right_indent is not None: pf.right_indent       = right_indent
    if first_line   is not None: pf.first_line_indent  = first_line
    return p


def insert_figure(doc, fname: str, caption: str, width=5.8):
    fig_path = FIG_DIR / fname
    if not fig_path.exists():
        print(f"  WARNING: figure not found: {fig_path}")
        return

    # Image row
    img_p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=12, space_after=3, lsp=1.0)
    img_p.paragraph_format.keep_with_next = True
    img_p.add_run().add_picture(str(fig_path), width=Inches(width))

    # Caption row
    cap_p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=3, space_after=14, lsp=1.15)
    parse_inline(cap_p, caption, size=10, italic=True)


def insert_table(doc, rows: list[list[str]], caption: str | None = None):
    if not rows:
        return
    # Table caption (above table, bold)
    if caption:
        cp = new_para(doc, space_before=10, space_after=3, lsp=1.15)
        parse_inline(cp, caption, size=BODY_SIZE - 1, bold=True)

    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_w = Inches(6.4 / ncols)

    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = tbl.rows[ri].cells[ci]
            cell.width = col_w
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            cell_text = row[ci].strip() if ci < len(row) else ""
            is_hdr = ri == 0
            parse_inline(p, cell_text, size=10, bold=is_hdr)
            if is_hdr:
                set_cell_bg(cell, "1E293B")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif ri % 2 == 0:
                set_cell_bg(cell, "F8FAFC")

    set_table_borders(tbl)
    # Spacer after table
    new_para(doc, space_before=2, space_after=8, lsp=1.0)


def insert_code_block(doc, code_lines: list[str]):
    if not code_lines:
        return
    p = new_para(doc, space_before=6, space_after=6, lsp=1.1,
                 left_indent=Inches(0.25), right_indent=Inches(0.25))
    set_para_shading(p, "F1F5F9")
    run = p.add_run("\n".join(code_lines))
    run.font.name = CODE_FONT
    run.font.size = Pt(9)


# ── Main document builder ─────────────────────────────────────────────────────

def build_docx():
    doc = Document()

    # Page setup — A4, 1-inch margins
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Inches(1)
    sec.right_margin  = Inches(1)
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(1)

    add_doc_header(doc, "MalariaLLM \u2014 Jemiri & Aladetola, 2026")
    add_page_number_footer(doc)

    # Base Normal style
    doc.styles["Normal"].font.name = BODY_FONT
    doc.styles["Normal"].font.size = Pt(BODY_SIZE)

    # ── Parse markdown ────────────────────────────────────────────────────────
    lines = MD_SRC.read_text(encoding="utf-8").splitlines()

    # State
    section  = "front"   # front | abstract | body | references
    in_code  = False
    code_buf = []
    tbl_rows = []
    tbl_caption_pending = None
    in_tbl = False
    after_author_block = False   # True when we just passed **Author:** / **Corresponding Author:**

    def flush_table():
        nonlocal in_tbl, tbl_rows, tbl_caption_pending
        if tbl_rows:
            insert_table(doc, tbl_rows, tbl_caption_pending)
        in_tbl = False
        tbl_rows = []
        tbl_caption_pending = None

    def flush_code():
        nonlocal code_buf
        insert_code_block(doc, code_buf)
        code_buf = []

    i = 0
    while i < len(lines):
        raw     = lines[i]
        stripped = raw.strip()

        # ── Code fences ──────────────────────────────────────────────────────
        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                if in_tbl:
                    flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # ── Table rows ───────────────────────────────────────────────────────
        if stripped.startswith("|"):
            if re.match(r'^\|[\s\-:|*]+\|', stripped):   # separator row
                i += 1
                continue
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            tbl_rows.append(cells)
            in_tbl = True
            i += 1
            continue
        else:
            if in_tbl:
                flush_table()

        # ── Blank line ───────────────────────────────────────────────────────
        if not stripped:
            after_author_block = False
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────────────────
        if stripped == "---":
            new_para(doc, space_before=4, space_after=4, lsp=1.0)
            i += 1
            continue

        # ── Detect section transitions from headings ─────────────────────────
        if stripped.startswith("## Abstract"):
            section = "abstract"
        elif stripped.startswith("## ") and section in ("front", "abstract"):
            section = "body"
        elif stripped.startswith("## References"):
            section = "references"

        # ── Headings ─────────────────────────────────────────────────────────
        if stripped.startswith("#### "):
            content = stripped[5:]
            p = new_para(doc, space_before=10, space_after=4, lsp=1.3)
            p.paragraph_format.keep_with_next = True
            parse_inline(p, content, size=11, bold=True, italic=True)
            i += 1
            continue

        if stripped.startswith("### "):
            content = stripped[4:]
            p = new_para(doc, space_before=12, space_after=4, lsp=1.3)
            p.paragraph_format.keep_with_next = True
            parse_inline(p, content, size=12, bold=True)
            i += 1
            continue

        if stripped.startswith("## "):
            content = stripped[3:]
            p = new_para(doc, space_before=16, space_after=6, lsp=1.3)
            p.paragraph_format.keep_with_next = True
            add_para_border_bottom(p)
            parse_inline(p, content, size=13, bold=True)
            i += 1
            continue

        if stripped.startswith("# "):
            content = stripped[2:]
            p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=0, space_after=10, lsp=1.25)
            parse_inline(p, content, size=16, bold=True)
            i += 1
            continue

        # ── Bullet list ──────────────────────────────────────────────────────
        if re.match(r'^[-*]\s', stripped):
            content = stripped[2:]
            indent_px = len(raw) - len(raw.lstrip())
            level = indent_px // 2
            p = new_para(doc, space_before=2, space_after=2, lsp=1.3,
                         left_indent=Inches(0.3 + level * 0.25),
                         first_line=Inches(-0.2))
            run = p.add_run("\u2022 ")
            run.font.name = BODY_FONT
            run.font.size = Pt(BODY_SIZE)
            parse_inline(p, content)
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────────
        nm = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if nm:
            num, content = nm.group(1), nm.group(2)
            if section == "references":
                p = new_para(doc, space_before=2, space_after=3, lsp=1.3,
                             left_indent=Inches(0.4), first_line=Inches(-0.4))
                run = p.add_run(f"{num}.\u00a0")
                run.font.name = BODY_FONT
                run.font.size = Pt(11)
                parse_inline(p, content, size=11)
            else:
                p = new_para(doc, space_before=2, space_after=3, lsp=1.3,
                             left_indent=Inches(0.4), first_line=Inches(-0.4))
                run = p.add_run(f"{num}.\u00a0")
                run.font.name = BODY_FONT
                run.font.size = Pt(BODY_SIZE)
                parse_inline(p, content)
            i += 1
            continue

        # ── Table caption line (**Table N. ...**) ────────────────────────────
        if re.match(r'^\*\*Table\s+\d+', stripped):
            inner = re.sub(r'^\*\*(.+?)\*\*$', r'\1', stripped.strip())
            tbl_caption_pending = inner if inner != stripped.strip() else stripped.strip("*")
            i += 1
            continue

        # ── Abstract section labels (**Background:** etc.) ───────────────────
        if re.match(r'^\*\*(Background|Objective|Methods|Results|Conclusions|Keywords):', stripped):
            p = new_para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_before=3, space_after=3, lsp=1.3,
                         left_indent=Inches(0.45), right_indent=Inches(0.45))
            parse_inline(p, stripped, size=11)
            i += 1
            continue

        # ── Front-matter author/metadata markers ─────────────────────────────
        if re.match(r'^\*\*(Running Title|Author|Corresponding Author):', stripped):
            after_author_block = True
            p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=5, space_after=2, lsp=1.2)
            parse_inline(p, stripped, size=11)
            i += 1
            continue

        # ── ORCID / Email lines ──────────────────────────────────────────────
        if stripped.startswith(("ORCID:", "Email:")):
            p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=1, space_after=1, lsp=1.2)
            parse_inline(p, stripped, size=10)
            i += 1
            continue

        # ── Plain text lines inside front-matter author block ────────────────
        # (author name, affiliation — appear right after **Author:** marker)
        if section == "front" and after_author_block:
            p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=1, space_after=1, lsp=1.2)
            parse_inline(p, stripped, size=11)
            i += 1
            continue

        # ── Regular body paragraph ────────────────────────────────────────────
        align = (WD_ALIGN_PARAGRAPH.JUSTIFY
                 if section in ("body", "abstract") else WD_ALIGN_PARAGRAPH.LEFT)
        p = new_para(doc, align=align,
                     space_before=0, space_after=6, lsp=1.5)
        parse_inline(p, stripped)

        # Figure injection: check trigger sentences
        for trigger, fname, caption in FIGURE_TRIGGERS:
            if trigger in stripped:
                insert_figure(doc, fname, caption)
                break

        i += 1

    # Final flush
    if in_tbl:
        flush_table()
    if in_code:
        flush_code()

    doc.save(str(OUT_DOC))
    size_kb = OUT_DOC.stat().st_size / 1024
    print(f"\n  DOCX saved : {OUT_DOC}")
    print(f"  Size       : {size_kb:.0f} KB\n")


if __name__ == "__main__":
    print("\nMalariaLLM \u2014 DOCX Generator")
    print("==============================\n")
    build_docx()
    print("  Done.\n")

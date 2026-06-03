"""
One-off error-corrections pass for the root manuscript DOCX
("final_paper_D (1).docx"):
  1. Section numbering -> consistent Arabic (matches publication/final_paper.md/.tex).
  2. Citation-mismatch flags -> visible red inline [CITATION CHECK: ...] notes
     appended in the same paragraph as each mismatched in-text citation.
     Original citation numbers are left unchanged (author to confirm sources).
  3. Dataset provenance wording -> "real-world" replaced with publicly-available
     Kaggle framing.
Run from repo root:  python publication/fix_root_docx.py
"""
import docx
from docx.shared import Pt, RGBColor

SRC = "final_paper_D (1).docx"
RED = RGBColor(0xC0, 0x00, 0x00)

# ── 1. Heading renumber map (exact current text -> new text) ──────────────────
HEADINGS = {
    "II. Related Work": "2. Related Work",
    "A. Rule-Based Expert Systems": "2.1 Rule-Based Expert Systems",
    "B. AI and Image-Based Malaria Diagnosis": "2.2 AI and Image-Based Malaria Diagnosis",
    "C. Large Language Models in Clinical Decision Support": "2.3 Large Language Models in Clinical Decision Support",
    "D. Explainable AI and Deployment Considerations": "2.4 Explainable AI and Deployment Considerations",
    "III. Methodology": "3. Methodology",
    "A. System Architecture": "3.1 System Architecture",
    "B. Dataset": "3.2 Dataset",
    "C. Rule-Based Classification": "3.3 Rule-Based Classification",
    "D. LLM Integration": "3.4 LLM Integration",
    "E. Evaluation Metrics": "3.5 Evaluation Metrics",
    "IV. Results": "4. Results",
    "Overall Performance": "4.1 Overall Performance",
    "B. Statistical Significance": "4.2 Statistical Significance",
    "C. Processing Efficiency": "4.3 Processing Efficiency",
    "D. LLM Explanation Quality": "4.4 LLM Explanation Quality",
    "E. Statistical Analysis": "4.5 Statistical Analysis",
    "1) McNemar’s Test": "4.5.1 McNemar’s Test",
    "2) Confidence Intervals": "4.5.2 Confidence Intervals",
    "3) Cohen’s Kappa": "4.5.3 Cohen’s Kappa",
    "4) Per-Class Statistical Testing": "4.5.4 Per-Class Statistical Testing",
    "5) Summary": "4.5.5 Summary",
    "V. Discussion": "5. Discussion",
    "VI. Conclusion and Future Work": "6. Conclusion and Future Work",
}

# ── 2. Citation flags: (anchor substring in paragraph, flag message) ──────────
F = " [CITATION CHECK: {}]"
FLAGS = [
    ("249 million cases worldwide",
     "in-text [10] cites the WHO disease-burden figures (249M cases / 608,000 deaths, 2022) but ref [10] is Cunningham et al. (RDT accuracy); the WHO World Malaria Report is ref [9]. ALSO in-text [23] cites the malaria economic/health-system burden but ref [23] is McNemar 1947 (a statistical test). Confirm intended sources."),
    ("microscopy and rapid diagnostic tests (RDTs) [2,11]",
     "in-text [2,11]: ref [2] (WHO Guidelines) is appropriate, but ref [11] is Garcez (neural-symbolic computing); a diagnostic-test/RDT reference (Cunningham, ref [10]) appears intended. Confirm."),
    ("assist health workers in low-resource settings [7,14]",
     "in-text [7,14]: ref [7] is Albahri (explainable AI) and ref [14] is Lievin (LLM reasoning); rule-based CDSS / expert-system references (Sutton ref [6]; Abu-Naser ref [13]) appear intended. Confirm."),
    ("structured and unstructured medical tasks [5,3,15]",
     "in-text [5,3,15]: ref [5] is Rajaraman (image detection) and ref [15] is Agarwal (mHealth); LLM clinical-reasoning references (e.g. Kung ref [4]; Lievin ref [14]) appear intended. Confirm."),
    ("patient privacy concerns when using cloud APIs [22,20]",
     "in-text [22,20]: ref [22] is the Kaggle dataset (Programmer3); intended references for hallucination/privacy risk appear to be Chen (ref [20]) and Wahl (ref [18]). Confirm."),
    ("infectious diseases [14]",
     "in-text [14] (expert-systems history: ONCOCIN/MYCIN) points to Lievin (LLM reasoning); an expert-systems reference (Abu-Naser, ref [13]) appears intended. Confirm."),
    ("rigid rule trees and symptom overlap [7]",
     "in-text [7] points to Albahri (explainable AI); a CDSS-limitations reference (Sutton, ref [6]) appears intended. Confirm."),
    ("high accuracy for parasite detection [6]",
     "in-text [6] points to Sutton (CDSS); a deep-learning parasite-detection reference (Rajaraman, ref [5]) appears intended. Confirm."),
    ("rural, resource-limited settings [11]",
     "in-text [11] points to Garcez (neural-symbolic computing); a diagnostic-test reference (Cunningham, ref [10]) appears intended. Confirm."),
    ("providing clinical recommendations [1,3,5,15]",
     "in-text [1,3,5,15]: refs [1] (Zhou) and [3] (Wang) fit, but ref [5] is Rajaraman (image detection) and ref [15] is Agarwal (mHealth); LLM references (Kung ref [4]; Lievin ref [14]) appear intended. Confirm."),
    ("underrepresented diseases or low-resource populations [22,21]",
     "in-text [22,21]: ref [22] is the Kaggle dataset and ref [21] is Ashley & Phyo (antimalarial drugs); references on LLM hallucination/bias (Chen ref [20]; Wang Y ref [19]) appear intended. Confirm."),
    ("clinician adoption of AI systems [8]",
     "in-text [8] points to Meta AI (Llama 3); an explainability-adoption reference (Albahri, ref [7]) appears intended. Confirm."),
    ("transparent explanations are essential [20,16,30]",
     "in-text [20,16,30]: ref [20] is Chen and ref [16] is Zhang; reference [30] DOES NOT EXIST (this list ends at [27]). References on offline AI/cost/transparency (Wahl ref [18]; Agarwal ref [15]; a privacy reference such as Kaissis, which is absent from this list) appear intended. Confirm and add the missing reference."),
    ("explanations referencing each patient’s symptoms [17]",
     "in-text [17] points to Marsh (life-threatening malaria); a prompt-engineering reference (Zhang, ref [16]) appears intended. Confirm."),
    ("cases from Kaggle Malaria Diagnosis Dataset [27]",
     "in-text [27] points to Dietterich (statistical tests); the Kaggle dataset is ref [22] (Programmer3). Confirm."),
    ("rules based on WHO guidelines [2,18]",
     "in-text [2,18]: ref [2] (WHO Guidelines) fits, but ref [18] is Wahl (AI in global health); a clinical-severity-indicators reference (Marsh, ref [17]) appears intended. Confirm."),
]


def set_heading_text(p, new_text):
    """Replace a heading paragraph's text, preserving the first run's formatting."""
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)


def replace_token(p, old, new):
    """Replace `old` with `new` inside a paragraph (single-run case, else rebuild)."""
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    joined = "".join(r.text for r in p.runs)
    if old in joined:
        new_joined = joined.replace(old, new)
        if p.runs:
            p.runs[0].text = new_joined
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(new_joined)
        return True
    return False


def append_flag(p, message):
    run = p.add_run(F.format(message))
    run.font.color.rgb = RED
    run.font.bold = True
    run.font.size = Pt(9)


def main():
    d = docx.Document(SRC)
    paras = d.paragraphs

    # 1. Headings
    done_head = set()
    for p in paras:
        t = p.text.strip()
        if t in HEADINGS and t not in done_head:
            set_heading_text(p, HEADINGS[t])
            done_head.add(t)
    missing_head = [h for h in HEADINGS if h not in done_head]
    print("Headings renumbered:", len(done_head), "/", len(HEADINGS))
    if missing_head:
        print("  !! NOT FOUND:", missing_head)

    # 3. Wording (do before flags so flag text isn't disturbed)
    w1 = w2 = False
    for p in paras:
        if "1,622 real-world, 60 synthetic" in p.text:
            w1 = replace_token(p, "1,622 real-world, 60 synthetic",
                               "1,622 from the publicly available Kaggle Malaria Diagnosis Dataset, 60 synthetic") or w1
        if p.text.strip().startswith("Real-world: 1,622 cases"):
            w2 = replace_token(p, "Real-world:", "Publicly available:") or w2
    print("Wording fix (abstract):", w1, "| (dataset bullet):", w2)

    # 2. Citation flags
    applied = 0
    for anchor, msg in FLAGS:
        hit = False
        for p in paras:
            if anchor in p.text:
                append_flag(p, msg)
                hit = True
                applied += 1
                break
        if not hit:
            print("  !! ANCHOR NOT FOUND:", anchor)
    print("Citation flags applied:", applied, "/", len(FLAGS))

    d.save(SRC)
    print("Saved", SRC)


if __name__ == "__main__":
    main()

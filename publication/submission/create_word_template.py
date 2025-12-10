"""
Create properly formatted Word document for JMIR submission - CORRECTED VERSION
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_jmir_manuscript():
    """Create JMIR-formatted manuscript template"""
    
    doc = Document()
    
    # Set margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Define styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    
    paragraph_format = normal_style.paragraph_format
    paragraph_format.line_spacing = 2.0
    paragraph_format.space_after = Pt(0)
    
    # ===== TITLE PAGE =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(
        'An LLM-Enhanced Clinical Decision Support System for Malaria Diagnosis '
        'in Resource-Limited Settings: A Hybrid Approach'
    )
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    
    doc.add_paragraph()
    
    # Running title
    running = doc.add_paragraph()
    running.alignment = WD_ALIGN_PARAGRAPH.CENTER
    running_run = running.add_run('Running Title: LLM-Enhanced Malaria Diagnosis System')
    running_run.font.name = 'Times New Roman'
    running_run.font.size = Pt(12)
    running_run.font.italic = True
    
    doc.add_paragraph()
    
    # Author information - CORRECTED
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author.add_run(
        'Daniel Jemiri\n'
        'PhD Student, Artificial Intelligence & Information Science\n'
        'The University of Texas at Dallas, Richardson, TX, United States\n'
        'ORCID: https://orcid.org/0009-0005-9338-7682'
    )
    author_run.font.name = 'Times New Roman'
    author_run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Corresponding author
    corresp = doc.add_paragraph()
    corresp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corresp_run = corresp.add_run(
        'Corresponding Author:\n'
        'Daniel Jemiri\n'
        'Email: Daniel.Jemiri@UTDallas.edu'
    )
    corresp_run.font.name = 'Times New Roman'
    corresp_run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ===== ABSTRACT =====
    abstract_heading = doc.add_heading('ABSTRACT', level=1)
    abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Background
    bg_heading = doc.add_paragraph()
    bg_run = bg_heading.add_run('Background: ')
    bg_run.bold = True
    bg_text = bg_heading.add_run(
        'Malaria remains a significant global health challenge, causing approximately '
        '600,000 deaths annually, predominantly in sub-Saharan Africa. Traditional rule-based '
        'expert systems for malaria diagnosis have limitations in handling symptom variability '
        'and providing clinical reasoning to support healthcare worker decisions.'
    )
    
    # Objective
    obj_heading = doc.add_paragraph()
    obj_run = obj_heading.add_run('Objective: ')
    obj_run.bold = True
    obj_text = obj_heading.add_run(
        'To develop and evaluate an LLM-enhanced clinical decision support system for malaria '
        'diagnosis that combines deterministic classification rules with large language model '
        'reasoning capabilities to improve diagnostic accuracy and explainability in '
        'resource-limited settings.'
    )
    
    # Methods
    method_heading = doc.add_paragraph()
    method_run = method_heading.add_run('Methods: ')
    method_run.bold = True
    method_text = method_heading.add_run(
        'We developed a hybrid system integrating a rule-based classifier with Llama 3.1 8B '
        'for natural language explanation generation. The rule-based component implements WHO '
        'malaria severity classification guidelines, while the LLM provides clinical reasoning, '
        'treatment recommendations, and confidence scoring. The system was evaluated on 1,682 '
        'malaria cases (1,622 real-world cases from the Kaggle Malaria Diagnosis Dataset plus '
        '60 synthetic cases) spanning four severity levels: No Malaria, Stage I (uncomplicated), '
        'Stage II (moderate), and Critical. Performance was compared against the original '
        'rule-based baseline using accuracy, precision, recall, and F1-score metrics.'
    )
    
    # Results
    results_heading = doc.add_paragraph()
    results_run = results_heading.add_run('Results: ')
    results_run.bold = True
    results_text = results_heading.add_run(
        'The LLM-enhanced hybrid system achieved 87.22% diagnostic accuracy (1,467/1,682 correct '
        'classifications) compared to 30.62% for the baseline rule-based system, representing a '
        '56.60 percentage point improvement (relative improvement of 185%). The system demonstrated '
        '100% accuracy (20/20) in detecting critical malaria cases, 98.9% accuracy (1,077/1,089) for '
        'Stage I cases, and 77.4% accuracy (352/455) for No Malaria cases. Average processing time '
        'was 6.3 seconds per diagnosis with an average LLM confidence score of 85%. The hybrid '
        'approach maintained the baseline\'s perfect critical case detection while dramatically '
        'improving Stage I detection from 1.8% to 98.9%.'
    )
    
    # Conclusions
    conclusions_heading = doc.add_paragraph()
    conclusions_run = conclusions_heading.add_run('Conclusions: ')
    conclusions_run.bold = True
    conclusions_text = conclusions_heading.add_run(
        'Integration of large language models with rule-based medical decision support systems '
        'significantly improves diagnostic accuracy while adding natural language explainability. '
        'This hybrid approach demonstrates promise for deployment in resource-limited healthcare '
        'settings where specialist physicians are scarce. The system\'s offline capability, zero '
        'API costs, and privacy-preserving architecture address key barriers to AI adoption in '
        'low- and middle-income countries. Future work should include prospective clinical '
        'validation and expansion to differential diagnosis of other febrile illnesses.'
    )
    
    doc.add_paragraph()
    
    # Keywords
    kw_heading = doc.add_paragraph()
    kw_run = kw_heading.add_run('Keywords: ')
    kw_run.bold = True
    kw_text = kw_heading.add_run(
        'malaria diagnosis; clinical decision support systems; large language models; '
        'artificial intelligence; resource-limited settings; explainable AI; hybrid systems; '
        'global health'
    )
    
    doc.add_page_break()
    
    # ===== INTRODUCTION =====
    doc.add_heading('INTRODUCTION', level=1)
    intro_placeholder = doc.add_paragraph(
        '[PASTE YOUR INTRODUCTION FROM results/paper_draft.md HERE]'
    )
    intro_placeholder.runs[0].font.italic = True
    intro_placeholder.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_page_break()
    
    # ===== METHODS =====
    doc.add_heading('METHODS', level=1)
    methods_placeholder = doc.add_paragraph(
        '[PASTE YOUR METHODOLOGY SECTION FROM results/paper_draft.md HERE]'
    )
    methods_placeholder.runs[0].font.italic = True
    methods_placeholder.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_page_break()
    
    # ===== RESULTS =====
    doc.add_heading('RESULTS', level=1)
    results_placeholder = doc.add_paragraph(
        '[PASTE YOUR RESULTS SECTION FROM results/paper_draft.md HERE]\n\n'
        '[INSERT TABLE 1: Overall Performance Metrics]\n\n'
        '[INSERT TABLE 2: Per-Stage Performance]\n\n'
        '[INSERT FIGURE 1: Confusion Matrices]\n\n'
        '[INSERT FIGURE 2: Accuracy Comparison Chart]\n\n'
        '[INSERT FIGURE 3: System Architecture Diagram]'
    )
    results_placeholder.runs[0].font.italic = True
    results_placeholder.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_page_break()
    
    # ===== DISCUSSION =====
    doc.add_heading('DISCUSSION', level=1)
    discussion_placeholder = doc.add_paragraph(
        '[PASTE YOUR DISCUSSION SECTION FROM results/paper_draft.md HERE]'
    )
    discussion_placeholder.runs[0].font.italic = True
    discussion_placeholder.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_page_break()
    
    # ===== CONCLUSION =====
    doc.add_heading('CONCLUSION', level=1)
    conclusion_placeholder = doc.add_paragraph(
        '[PASTE YOUR CONCLUSION FROM results/paper_draft.md HERE]'
    )
    conclusion_placeholder.runs[0].font.italic = True
    conclusion_placeholder.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_page_break()
    
    # ===== DATA AVAILABILITY - CORRECTED =====
    doc.add_heading('Data Availability Statement', level=1)
    data_avail = doc.add_paragraph(
        'The dataset used in this study is publicly available:\n'
        '• Synthetic dataset: Available at https://github.com/jemiridaniel/malaria-llm-cdss/data\n'
        '• Real-world data: Kaggle Malaria Diagnosis Dataset '
        '(https://kaggle.com/datasets/programmer3/malaria-diagnosis-dataset)\n'
        '• Source code: https://github.com/jemiridaniel/malaria-llm-cdss\n'
        '• Evaluation results: Included in supplementary materials\n\n'
        'All code, data processing scripts, and evaluation notebooks are openly available under '
        'the MIT License to facilitate reproducibility and future research.'
    )
    
    # ===== ETHICS STATEMENT =====
    doc.add_heading('Ethics Statement', level=1)
    ethics = doc.add_paragraph(
        'This computational study utilized publicly available datasets and synthetic data. '
        'No human subjects were directly involved in data collection. The study does not require '
        'Institutional Review Board (IRB) approval as it constitutes a retrospective analysis of '
        'de-identified datasets.\n\n'
        'The system is designed and presented as clinical decision support, not autonomous diagnosis. '
        'Implementation includes appropriate safety warnings and human-in-the-loop requirements for '
        'clinical deployment. Future prospective clinical validation studies will require appropriate '
        'ethical approvals.'
    )
    
    # ===== ACKNOWLEDGMENTS - CORRECTED =====
    doc.add_heading('Acknowledgments', level=1)
    ack = doc.add_paragraph(
        'We acknowledge the open-source community for Ollama and Meta AI for Llama 3.1. We thank '
        'the contributors to the Kaggle Malaria Diagnosis Dataset for making their data publicly '
        'available. The original undergraduate system was developed at Federal University of '
        'Technology Owerri (FUTO), Imo State, Nigeria, in 2020 under the supervision of '
        'Dr. Stanley Obilor, B.Sc. (Nig.), M.Sc (FUTO), Ph.D (Nig), MNSE, MNICE, MNIWE.'
    )
    
    # ===== CONFLICTS =====
    doc.add_heading('Conflicts of Interest', level=1)
    conflicts = doc.add_paragraph(
        'None declared. This research was conducted independently without external funding or '
        'commercial interests.'
    )
    
    doc.add_page_break()
    
    # ===== REFERENCES =====
    doc.add_heading('REFERENCES', level=1)
    ref_placeholder = doc.add_paragraph(
        '[PASTE THE 30 FORMATTED REFERENCES FROM publication/references/references_AMA_formatted.txt HERE]'
    )
    ref_placeholder.runs[0].font.italic = True
    ref_placeholder.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    # Save document
    output_path = 'publication/submission/Jemiri_Malaria_LLM_CDSS_Manuscript_CORRECTED.docx'
    doc.save(output_path)
    print(f"✅ CORRECTED Word document created: {output_path}")
    print("\nCorrections made:")
    print("  ✅ FUTA → FUTO (Federal University of Technology Owerri, Imo State)")
    print("  ✅ Removed MS degree")
    print("  ✅ Added Dr. Stanley Obilor as undergraduate supervisor")
    print("  ✅ Clarified: First-year PhD student at UTD")
    print("\nNext steps:")
    print("1. Open the document in Word")
    print("2. Replace all RED TEXT with your actual content")
    print("3. Insert tables and figures")
    print("4. Add page numbers")
    print("5. Review and save!")

if __name__ == "__main__":
    create_jmir_manuscript()

